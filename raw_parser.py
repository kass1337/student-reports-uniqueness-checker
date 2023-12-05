import docx
import error_codes
class raw_parser(object):
    doc = None
    _default_heading = "None-Heading"
    _default_table = "Table "

    def __init__(self):
        pass

    def raw_parse(self, file_name):
        try:
            self.doc = docx.Document(file_name)
        except docx.opc.exceptions.PackageNotFoundError:
            exit(error_codes.FILE_NOT_FOUND.value)
        current_heading = self._default_heading
        content = {}
        heading_content = {}
        content[current_heading] = list()
        for par in self.doc.paragraphs:
            if "Heading" in par.style.name:
                current_heading = par.text
                content[current_heading] = list()
                continue
            content[current_heading].append(par.text)
        for i, table in enumerate(self.doc.tables):
            content[self._default_table + str(i)] = list()
            for j, row in enumerate(table.rows):
                text = list(cell.text for cell in row.cells)
                content[self._default_table + str(i)].extend(text)
        self._split_headings(content, heading_content)
        return content, heading_content
    def _split_headings(self, content, heading_content):
        cur_heading_full_text = " "
        for key in content:
            for part in content[key]:
                cur_heading_full_text += part + " "
            heading_content[key] = cur_heading_full_text
            cur_heading_full_text = " "